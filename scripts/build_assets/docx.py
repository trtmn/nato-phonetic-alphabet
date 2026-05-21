"""DOCX generator: matches the PDF's two-column layout where DOCX allows."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from nato_phonetic.core import NATO_PHONETIC_ALPHABET

from . import config


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shading)


def build_docx(dest: Path) -> None:
    """Generate a printable DOCX of the NATO phonetic alphabet."""
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(config.TITLE)
    run.font.name = "Source Code Pro"
    run.font.size = Pt(config.HEADER_FONT_SIZE)

    doc.add_paragraph()

    letters = [chr(c) for c in range(ord("A"), ord("Z") + 1)]
    table = doc.add_table(rows=13, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i in range(13):
        left, right = letters[i], letters[i + 13]
        row = table.rows[i]
        cells = row.cells
        values = [left, NATO_PHONETIC_ALPHABET[left], right, NATO_PHONETIC_ALPHABET[right]]
        for col_idx, value in enumerate(values):
            cell = cells[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if col_idx % 2 == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            run = paragraph.add_run(value)
            run.font.name = "Source Code Pro"
            run.font.size = Pt(config.BODY_FONT_SIZE)
            if i % 2 == 1:
                _set_cell_shading(cell, config.ZEBRA_COLOR)

    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f"{config.TITLE} {config.LICENSE_TEXT}")
    footer_run.font.name = "Source Code Pro"
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.save(str(dest))
